# this is 90% "borrowed code"
# https://9to5answer.com/adding-an-hyperlink-in-msword-by-using-python-docx

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.add_break()

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


document = docx.Document()
p = document.add_paragraph()

cve_array = [] # Provide a list of CVE numbers
#EXAMPLE cve_array = ['CVE-2015-6161','CVE-2016-0199','CVE-2016-0200']
BASEURL = 'https://nvd.nist.gov/vuln/detail/' # this is the baseurl for the CVE links.

for cve in cve_array:
    add_hyperlink(p, cve, BASEURL + cve)
document.save('demo_hyperlink.docx')
